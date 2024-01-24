import streamlit as st
import docx
import llm_agent
import openai
from datetime import datetime
from langchain.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain.chat_models import ChatOpenAI
from langchain.agents import load_tools
from langchain.agents import initialize_agent
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from dotenv import load_dotenv
from langchain.callbacks import get_openai_callback
import datetime
from langchain.chains import RetrievalQA
import io
import base64
from docx import Document
import os
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import numpy as np
import pandas as pd
import csv
import inspect
from typing import List
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.pydantic_v1 import BaseModel, Field
import llm_tools_agent
from docx.shared import Inches


load_dotenv()


os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
os.environ["OPENAI_API_KEY"] = st.secrets["SERPAPI_API_KEY"]
# os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')
# os.environ['SERPAPI_API_KEY'] = os.getenv['SERPAPI_API_KEY'] 


chat_llm = ChatOpenAI(model='gpt-3.5-turbo',temperature=0.0)

embedding = OpenAIEmbeddings()
tool_names = ["serpapi"]
tools = load_tools(tool_names)

agent = initialize_agent(tools, chat_llm, agent="zero-shot-react-description", verbose=True,handle_parsing_errors=True)


class step1_risk_struture(BaseModel):
        risk_category: str = Field(description="category of the risk")
        risk_title: str = Field(description="title of the risk")
        risk_description: str = Field(description="description of the risk")
        severity: str = Field(description="severity of the risk")
        likelihood: str = Field(description="likelihood of the risk")
        Preliminary_Risk_Level: str = Field(description="Preliminary Risk Level of the risk")

class step2_risk_rank_struture(BaseModel):
        risk_category: str = Field(description="category of the risk")
        risk_title: str = Field(description="title of the risk")
        risk_description: str = Field(description="description of the risk")
        severity: str = Field(description="severity of the risk")
        likelihood: str = Field(description="likelihood of the risk")
        Preliminary_Risk_Level: str = Field(description="Preliminary Risk Level of the risk")
        Overall_Risk_Level: str = Field(description="Overall Risk Level of the risk")

class actionable_struture(BaseModel):
        risk_category: str = Field(description="category of the risk")
        risk_title: str = Field(description="title of the risk")
        risk_description: str = Field(description="description of the risk")
        severity: str = Field(description="severity of the risk")
        likelihood: str = Field(description="likelihood of the risk")
        Preliminary_Risk_Level: str = Field(description="Preliminary Risk Level of the risk")
        Overall_Risk_Level: str = Field(description="Overall Risk Level of the risk")
        Actionables: str = Field(description="Actionables of the risk")

def get_ranking_priority(use_case_summary,step1_results):
    
    parser = JsonOutputParser(pydantic_object=step2_risk_rank_struture)

    prompt = PromptTemplate(
        template="""
        List only Overall high and medium risk using the Following.
        Use Case Summary : {use_case_summary} 
        Risk Information : {step1_results}
        \n{format_instructions}\n""",
        input_variables=["use_case_summary","step1_results"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    chain = prompt | chat_llm | parser

    step2_result=chain.invoke({"use_case_summary": use_case_summary,"step1_results":step1_results})
    return step2_result


def generate_risk_information(use_case_summary,risk_title):
    
    parser = JsonOutputParser(pydantic_object=step1_risk_struture)

    prompt = PromptTemplate(
        template="""
                Provide the Risk description based on risk title and use case summary.
                Rank the risk associated with the use case based on severity and likelihood as high medium and low. Also Provide a Preliminary risk score for every risk.
                Risk Title : {risk_title}
                Use Case Summary :\n{query}\n{format_instructions}\n""",
        input_variables=["query","risk_title"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    chain = prompt | chat_llm | parser

    step1_result=chain.invoke({"query": use_case_summary,"risk_title":risk_title})
    return step1_result

def generate_risk_information_tools(use_case_summary,risk_title,search_result):
    
    parser = JsonOutputParser(pydantic_object=step1_risk_struture)

    prompt = PromptTemplate(
        template="""
                Provide the Risk description based on risk title , context and use case summary.
                Rank the risk associated with the use case based on severity and likelihood as high medium and low. Also Provide a Preliminary risk score [High/Medium/Low] for every risk.
                Risk Title : {risk_title}
                Context : {search_result}
                Use Case Summary :\n{query}\n{format_instructions}\n""",
        input_variables=["query","risk_title","search_result"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    chain = prompt | chat_llm | parser

    step1_result=chain.invoke({"query": use_case_summary,"risk_title":risk_title,"search_result":search_result})
    return step1_result

def generate_actionable(use_case_summary,risk_info):
    
    for attempt in range(5):
        try:
            parser = JsonOutputParser(pydantic_object=actionable_struture)

            prompt = PromptTemplate(
                template="""
                Provide Actionable steps for governance to address each identified risk for Given risk Information and Use Case Summary.
                For each risk compile a set of actionables to address the risk. These actionables shall be governance actionables.
                Use Case Summary : {use_case_summary} 
                Risk Information : {risk_info}
                \n{format_instructions}\n""",
                input_variables=["use_case_summary","risk_info"],
                partial_variables={"format_instructions": parser.get_format_instructions()},
            )

            chain = prompt | chat_llm | parser

            actionbles=chain.invoke({"use_case_summary": use_case_summary,"risk_info":risk_info})
            return actionbles
        except Exception as e:
            if attempt < 4:  # 0-indexed, so 4 is the 5th attempt
                continue  # Try again
            else:
                print(f"An error occurred after 5 attempts: {e}")
                return None


def read_docx(file_path):
    """
    Read the content of a docx file and return it as a text string.
    """
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def save_docx(content, filename):
    # Create a new Document object
    doc = Document()

    # Add content to the document
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)

    # Save the document
    doc.save(filename)

# Function to create a download link
def create_download_link(filename):
    with open(filename, "rb") as file:
        data = file.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:file/docx;base64,{b64}" download="{filename}">Download Risk Assessment Document</a>'
    return href




def save_to_excel(use_case_title, use_case_summary, risk_ranking_df, actionables_df, filename):
    with pd.ExcelWriter(filename) as writer:
        # Use Case Title and Summary
        df_summary = pd.DataFrame({'Use Case Title': [use_case_title], 'Use Case Summary': [use_case_summary]})
        df_summary.to_excel(writer, sheet_name='Summary', index=False)

        # Risk Ranking
        risk_ranking_df.to_excel(writer, sheet_name='Risk Ranking', index=False)

        # Actionables
        actionables_df.to_excel(writer, sheet_name='Actionables', index=False)



def generate_use_case_summary(df,use_case_title):
    qa_list = df.set_index('Questions').to_dict()['Answer']
    title_template = """
    Create an LLM downstream use case context Description as a paragraph for "{use_case_title}"with all of the following information given in list.
         Information:"{qa_list}"
                """
    prompt = ChatPromptTemplate.from_template(template=title_template)
    messages = prompt.format_messages(qa_list=qa_list,use_case_title=use_case_title)
    response = chat_llm(messages)
    use_case_summary = response.content
    
    return use_case_summary

# def generate_risks(use_case_summary):
    
#     file_path = r'Documented nature of risks.docx'
#     documented_nature_of_risks = read_docx(file_path)
    
#     title_template = """
#         Identify the risks that apply to the use case. Use the information in Knowledge base to identify the applicable risks. 
#         Provide atleast 1 or 2 examples of each risk using the use case brief and the user responses to the questions.
#         Knowledge base: {context}
#         use case: {question}
#                 """
#     prompt = ChatPromptTemplate.from_template(template=title_template)
#     messages = prompt.format_messages(question=use_case_summary, context=documented_nature_of_risks)
#     response = chat_llm(messages)
#     risk_information = response.content
#     ps2_doc.add_paragraph("Key Risks:")
#     ps2_doc.add_paragraph(risk_information)
#     return risk_information

def generate_risks(use_case_summary):
    
##################################### 1) agents Run############################################################

    # Initialize an empty list to store the results
    results_list = []

    # Iterate over all functions in llm_agent
    for name, func in inspect.getmembers(llm_agent, inspect.isfunction):
        try:
            # Call the function and get the result
            agent_result = func(use_case_summary)
            print(agent_result)

            # Check if the result is a dictionary or a list of dictionaries
            if isinstance(agent_result, dict):
                # Append the result to the list
                results_list.append(agent_result)
            elif isinstance(agent_result, list) and all(isinstance(item, dict) for item in agent_result):
                # Extend the list with multiple dictionaries
                results_list.extend(agent_result)
        except Exception as e:
            # Handle any errors that occur during function execution
            print(f"Error in {name}: {e}")

    # Write the results to a CSV file
    if results_list:
        keys = results_list[0].keys()
        with open('output.csv', 'w', newline='') as output_file:
            dict_writer = csv.DictWriter(output_file, keys)
            dict_writer.writeheader()
            dict_writer.writerows(results_list)
    else:
        print("No valid results to write to CSV.")
        
###################################### 2) agents Step 1#############################################################
    df1= pd.read_csv('output.csv')
    final_df = pd.DataFrame(columns=['risk_category', 'risk_title', 'risk_description', 'severity', 'likelihood', 'Preliminary_Risk_Level'])

    for index, row in df1.iterrows():
        risk_title = row['risk_title']
        
        result = generate_risk_information(use_case_summary, risk_title)

        # Append the result to the final DataFrame
        final_df = final_df.append(result, ignore_index=True)
        
    # final_df.to_csv('final_output.csv', index=False)  
###################################### 3) Google Search ############################################################
        
    # doc = Document()

    # user_response = ["gte-small", "llama2", "word2vec", "faiss", "prompt chain"]
    # heading = ["embedding model", "LLM", "embedding approach", "vector database", "prompting approach"]

    # model_dict = dict(zip(user_response, heading))

    # aspects = ["limitations", "constraints", "disadvantages", "issues", "risks"]

    # for model, head in model_dict.items():
    #     tool_response = ""
    #     for aspect in aspects:
    #         prompt = f"{aspect} of using {model} {head} in 3-4 bullet points."
    #         print(prompt)
            
    #         response = agent.run(prompt)
    #         tool_response += response + "\n" 
        
    #     doc.add_heading(model, level=1)
    #     doc.add_paragraph(tool_response) 
    # doc.save("gs_results.docx")


###################################### 4) Agent_tool Run on Search results and use case #############################


    search_result=read_docx("gs_results.docx")
    # Initialize an empty list to store the results
    results_list = []

    # Iterate over all functions in llm_agent
    for name, func in inspect.getmembers(llm_tools_agent, inspect.isfunction):
        try:
            # Call the function and get the result
            agent_result = func(use_case_summary=use_case_summary,search_result=search_result)
            print(agent_result)

            # Check if the result is a dictionary or a list of dictionaries
            if isinstance(agent_result, dict):
                # Append the result to the list
                results_list.append(agent_result)
            elif isinstance(agent_result, list) and all(isinstance(item, dict) for item in agent_result):
                # Extend the list with multiple dictionaries
                results_list.extend(agent_result)
        except Exception as e:
            # Handle any errors that occur during function execution
            print(f"Error in {name}: {e}")

    # Write the results to a CSV file
    if results_list:
        keys = results_list[0].keys()
        with open('output2.csv', 'w', newline='') as output_file:
            dict_writer = csv.DictWriter(output_file, keys)
            dict_writer.writeheader()
            dict_writer.writerows(results_list)
    else:
        print("No valid results to write to CSV.")
        
        
###################################### 5)Agent tool result with step 1 ############################################
    df2=pd.read_csv('output2.csv')
    final_df2 = pd.DataFrame(columns=['risk_category', 'risk_title', 'risk_description', 'severity', 'likelihood', 'Preliminary_Risk_Level'])

    for index, row in df2.iterrows():
        risk_title = row['risk_title']
        
        result = generate_risk_information_tools(use_case_summary, risk_title,search_result)

        # Append the result to the final DataFrame
        final_df2 = final_df2.append(result, ignore_index=True)

    # Print the final DataFrame
    # final_df2.to_csv('final_output2.csv', index=False) 
    
###################################### 6) Combine and Filter and Show ##############################################
    # df3 = pd.read_csv('Final_output.csv')
    # df4 = pd.read_csv('Final_output2.csv')

    # Concatenate the dataframes
    df = pd.concat([final_df, final_df2])

    # Filter rows where 'Preliminary risk score' is 'High' or 'Medium'
    df = df[df['Preliminary_Risk_Level'].isin(['High', 'Medium'])]

    # Save the new dataframe
    # df.to_csv('final_df.csv', index=False)
    
    return df
    




def rank_risks(use_case_summary,df):
    final_df3 = pd.DataFrame(columns=['risk_category', 'risk_title', 'risk_description', 'severity', 'likelihood', 'Preliminary_Risk_Level','Overall_Risk_Level'])

    for index, row in df.iterrows():
        # Convert the row to a dictionary and print it
        result=get_ranking_priority(use_case_summary=use_case_summary,step1_results=row.to_dict())
        
        final_df3 = final_df3.append(result, ignore_index=True)

    # Print the final DataFrame
    # final_df3.to_csv('all_risk_rank_info.csv', index=False)
    
    return final_df3
    
    

def mitigate_risks(risk_df,use_case_summary):
    st.subheader("Actionables for Risk Mitigation")
    # Iterate over each row in the dataframe
    final_df4 = pd.DataFrame(columns=['risk_category', 'risk_title', 'risk_description', 'severity', 'likelihood', 'Preliminary_Risk_Level','Overall_Risk_Level','Actionables'])

    for index, row in risk_df.iterrows():
        # Convert the row to a dictionary and print it
        result=generate_actionable(use_case_summary=use_case_summary,risk_info=row.to_dict())
        
        final_df4 = final_df4.append(result, ignore_index=True)
        
    
    final_df4 = final_df4[['risk_description','Actionables']]
    # Print the final DataFrame
    # final_df4.to_csv('actionables_with risk.csv', index=False)
    
    return final_df4

# def generate_summary(doc):
#     st.subheader("Final Summary")
#     summary = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    
#     title_template = """Compile All information in "{summary}". Structure in the following format:
#         The document shall contain the following information: 
#         Section A: Brief about the use case. 
#         Section B: List of high-level risks associated with the use case.
#         Section C: Table containing key risks with their risk ranking along with the reasons for the risk ranking.
#         Section D: List of actionables for each risk listed in Section C.
#     """
#     prompt = ChatPromptTemplate.from_template(template=title_template)
#     messages = prompt.format_messages(summary=summary)
#     response = chat_llm(messages)
#     final_document = response.content
#     return final_document
   
#UI

def main():
        
    st.title("LLM Risk Assessment Engine")

    # Print the provided text
    st.write("""
    The LLM Risk Assessment Engine is designed to specialize in risk evaluation, providing insights and assessments for a range of risks including content risks, context risks, trust risks, and societal and sustainability risks.
    """)
    # Step 1: Get the use case title from the user
    use_case_title = st.text_input("Write Your Use Case")

    # Step 2: Ask Questions
    # Step 3 : Get Answer
    if use_case_title:
        st.write(f"""Provide a brief about the {use_case_title}.""")
        questions = [
        'What is the industry of the use case?',
        'What is the nature of the use case?',
        'What type of data is used?',
        'What Language Model (LLM) is used?',
        'What embedding approach is used?',
        'What embedding model is used?',
        'Who is the user group?',
        'Is fine-tuning applied?',
        'What is the domain context?',
        'What vector store is used?',
        'Is prompt engineering applied?',
        'How is the model deployed?',
        'How is the model monitored?',
        'What logging and feedback mechanisms are used?',
        'What guardrails are applied?'
        ]
        df = pd.DataFrame(questions, columns=['Questions'])
        df['Answer'] = ''

        config = {
            'Answer' : st.column_config.TextColumn('Answer (required)', width='large', required=True),
        }

        result = st.data_editor(df, column_config = config, num_rows='dynamic')
        result.to_csv('questions_and_answers.csv', index=False)
        
        if st.button('Get results'):
            st.write(result)
            st.session_state['result'] = result  # Store result in session state

        if 'result' in st.session_state:
            if st.session_state['result']['Answer'].isnull().any():
                st.write("Please answer all questions before continuing.")
            elif st.button("Continue"):
                with st.spinner('Processing your use_cases...'):
                    data=pd.read_csv('questions_and_answers.csv')
                    use_case_summary = generate_use_case_summary(data,use_case_title)
                    st.session_state['use_case_summary'] = use_case_summary  # Store use_case_summary in session state
                    st.write(use_case_summary)

        if 'use_case_summary' in st.session_state:
            if st.button("Generate Risk Assessment Document"):
                with st.spinner('Generating Risks...'):
                    risk_information = generate_risks(st.session_state['use_case_summary'])  # Use use_case_summary from session state
                    st.session_state['risk_information'] = risk_information  # Store risk_information in session state
                    st.subheader("Risk with Examples")
                    st.write(risk_information)

                with st.spinner('Ranking risks...'):
                    risk_ranking = rank_risks(df=st.session_state['risk_information'], use_case_summary=st.session_state['use_case_summary'])  # Use risk_information from session state
                    st.session_state['risk_ranking'] = risk_ranking  # Store risk_ranking in session state
                    st.subheader("Risk Ranking")
                    st.write(risk_ranking)

                with st.spinner('Generating actionables for risk mitigation...'):
                    actionables = mitigate_risks(risk_df=st.session_state['risk_ranking'], use_case_summary=st.session_state['use_case_summary'])  # Use risk_ranking from session state
                    st.session_state['actionables'] = actionables  # Store actionables in session state
                    st.write(actionables)

                with st.spinner('Compiling the final document...'):
                    # doc=read_docx("gs_results.docx")
                    filename = f"{use_case_title} Risk Assessment.xlsx" 
                    save_to_excel(use_case_title=use_case_title, use_case_summary=st.session_state['use_case_summary'], risk_ranking_df=st.session_state['risk_ranking'], actionables_df=st.session_state['actionables'], filename=filename)

                    
                    download_link = create_download_link(filename)
                    st.session_state['download_link'] = download_link  # Store download_link in session state
                    st.markdown(st.session_state['download_link'], unsafe_allow_html=True)  # Use download_link from session state
                    



if __name__ == "__main__":
    main()

